from collections import defaultdict

from docx import Document

from pur_doc import constant, sql


class myDefaultDict(defaultdict):
    __repr__ = dict.__repr__


def fake_dict():
    return myDefaultDict(fake_dict)


def inject_docx(project_id, vendor_id, part_id_list):
    document = Document()

    # prepare data

    # inject project information
    project = sql.get_project_info(project_id)
    if project is None:
        project = fake_dict()

    vendor = sql.get_vendor_info(vendor_id)
    if vendor is None:
        vendor = fake_dict()

    part_list = []
    for part_id in part_id_list:
        part = sql.get_part_general_info(project_id, part_id)
        if part:
            part_list.append(part)

    document.add_heading(f'Nomination Letter for project {project}')

    for part in part_list:
        document.add_heading(f'{part["part"]} {part["part_description"]}')

    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    document.add_heading('The role of dolphins', level=2)

    table = document.add_table(rows=2, cols=2)

    cell = table.cell(0, 1)

    cell.text = 'parrot, possibly dead'

    document.add_paragraph('Lorem ipsum dolor sit amet.', style='ListBullet')

    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    paragraph.style = 'List Bullet'

    document.save('test.docx')
